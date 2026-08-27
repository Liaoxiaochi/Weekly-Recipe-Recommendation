"""Insert Chapters 3, 4 and 5 into the dissertation and apply the Ch1/Ch2 fixes.

Reads   毕业论文_修订版.docx   (never modified)
Writes  毕业论文_v5.docx

Styles are set by writing the template's style IDs directly rather than by
name, because the Leeds template uses numeric IDs ('1', '2', '3') for the
heading styles and python-docx resolves styles by name.

The Word table of contents is a field and cannot be refreshed from here.
Open the result in Word and press Ctrl+A then F9 to update it.
"""

import os
import shutil

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ch3_content import (BLOCKS as CH3_BLOCKS, CH1_USDA_FIXES,
                         CH2_CAPTION_FIXES, CH2_CROSSREF_FIXES,
                         NEW_REFERENCES)
from ch4_content import BLOCKS as CH4_BLOCKS
from ch5_content import BLOCKS as CH5_BLOCKS
from ch6_content import BLOCKS as CH6_BLOCKS
from frontmatter_content import (REPLACEMENTS as FM_REPLACEMENTS,
                                 DELETE_CONTAINING as FM_DELETE,
                                 REPLACE_WITH_BLOCKS as FM_BLOCKS,
                                 DELIVERABLES, DELIVERABLES_HEADER)
from appendix_content import (APPENDIX_A, APPENDIX_B,
                              REPLACE_FROM_HEADING)
from ch1_additions import (INSERT_AFTER as CH1_INSERTS,
                           REPLACEMENTS as CH1_REPLACEMENTS,
                           NUMBERED_LISTS)
from ch2_content import (INSERTIONS as CH2_INSERTIONS,
                         REPLACEMENTS as CH2_REPLACEMENTS,
                         DELETE_CONTAINING as CH2_DELETE,
                         TABLE_2_1_ADDITIONS,
                         NEW_REFERENCES as CH2_REFERENCES)

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, "..")
SRC = os.path.join(ROOT, "毕业论文_修订版.docx")
DST = os.path.join(ROOT, "毕业论文_v5.docx")
# A copy under a name that never changes.  The versioned file is for
# rolling back; this one is what any instruction to the author should
# name, so that bumping the version number cannot silently send them to
# a stale draft -- which is exactly what happened between v4 and v5.
LATEST = os.path.join(ROOT, "毕业论文_最新.docx")
FIG = os.path.join(ROOT, "figures")

report = []


def log(msg):
    report.append(msg)
    print(msg)


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def set_style_id(par, style_id):
    """Apply a style by its w:styleId, bypassing python-docx's name lookup."""
    pPr = par._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pStyle")):
        pPr.remove(old)
    el = OxmlElement("w:pStyle")
    el.set(qn("w:val"), style_id)
    pPr.insert(0, el)


def para_text(par):
    return "".join(r.text for r in par.runs)


def rewrite_paragraph_text(par, new_text):
    """Replace a paragraph's text, keeping the formatting of its first run.

    Word splits a sentence across many runs, so a target substring often does
    not exist contiguously in any single run.  Collapsing to one run is safe
    for these plain body paragraphs, none of which carries inline formatting.
    """
    runs = par.runs
    if not runs:
        par.add_run(new_text)
        return
    keep = runs[0]
    keep.text = new_text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)


# ---------------------------------------------------------------------------
# open
# ---------------------------------------------------------------------------
shutil.copyfile(SRC, DST)
doc = docx.Document(DST)
body = doc.element.body

# locate the List of References heading -- Chapter 3 goes immediately before it
ref_heading = None
for p in doc.paragraphs:
    if para_text(p).strip() == "List of References":
        ref_heading = p
        break
if ref_heading is None:
    raise SystemExit("could not find the 'List of References' heading")
log(f"anchor found: 'List of References'")


# ---------------------------------------------------------------------------
# Chapter 2 fix 1 -- swap the inverted table numbers
# ---------------------------------------------------------------------------
log("\n--- Chapter 2: table numbering ---")
pending = list(CH2_CAPTION_FIXES)
for p in doc.paragraphs:
    t = para_text(p).strip()
    for old, new in list(pending):
        if t == old.strip():
            rewrite_paragraph_text(p, new)
            log(f"  renumbered: {new[:46]}...")
            pending.remove((old, new))
            break
if pending:
    raise SystemExit(f"caption not matched: {pending}")


# ---------------------------------------------------------------------------
# Chapter 2 fix 2 -- add the missing in-text references to Figure 2.1,
# Table 2.1 and Table 2.2
# ---------------------------------------------------------------------------
log("\n--- Chapter 2: cross-references ---")
pending = list(CH2_CROSSREF_FIXES)
for p in doc.paragraphs:
    t = para_text(p)
    for anchor, addition in list(pending):
        if anchor in t:
            rewrite_paragraph_text(p, t.rstrip() + addition)
            log(f"  added:{addition[:58]}...")
            pending.remove((anchor, addition))
            break
if pending:
    raise SystemExit(f"cross-reference anchor not matched: {pending}")


# ---------------------------------------------------------------------------
# Chapter 1 fix -- remove the USDA promise superseded by design decision DD-2
# ---------------------------------------------------------------------------
log("")
log("--- Chapter 1: corrections and additions to Section 1.4 ---")
n_rep = 0
for _p in doc.paragraphs:
    _t = para_text(_p)
    _new = _t
    for _old, _rep in CH1_REPLACEMENTS:
        _new = _new.replace(_old, _rep)
    if _new != _t:
        rewrite_paragraph_text(_p, _new)
        n_rep += 1
        log("  corrected: ..." + _new[:64] + "...")
log("  " + str(n_rep) + " paragraph(s) corrected")

n_ins = 0
for _anchor, _additions in CH1_INSERTS:
    _target = None
    for _p in doc.paragraphs:
        if _anchor in para_text(_p):
            _target = _p
            break
    if _target is None:
        raise SystemExit("Chapter 1 anchor not found: " + _anchor)
    for _extra in reversed(_additions):
        _np = doc.add_paragraph(_extra)
        set_style_id(_np, "a")
        _target._p.addnext(_np._p)
        n_ins += 1
    log("  inserted " + str(len(_additions)) + " paragraph(s)")
log("  " + str(n_ins) + " paragraph(s) added to Section 1.4")

log("\n--- Chapter 1: USDA wording ---")
n_usda = 0
for p in doc.paragraphs:
    t = para_text(p)
    if "USDA" not in t:
        continue
    new = t
    new = new.replace(
        "per-recipe nutritional values from USDA FoodData Central and "
        "ingredient-level allergen tags",
        "per-serving nutritional values normalised from the corpus's own "
        "nutrition fields and ingredient-level allergen tags")
    new = new.replace(
        "allergen tags and per-recipe nutritional values from USDA "
        "FoodData Central.",
        "allergen tags and per-serving nutritional values normalised to "
        "absolute quantities.")
    new = new.replace("publicly licensed sources", "a single public dataset")
    new = new.replace(" and USDA FoodData Central", "")
    if new != t:
        rewrite_paragraph_text(p, new)
        n_usda += 1
        log(f"  reworded: ...{new[:70]}...")
log(f"  {n_usda} paragraph(s) changed")
for p in doc.paragraphs:
    if "USDA" in para_text(p):
        raise SystemExit("a USDA mention survived: " + para_text(p)[:90])


# ---------------------------------------------------------------------------
# Chapter 1 -- label and indent the objectives and deliverables
#
# Runs after the rewording above so that a paragraph is labelled once, in its
# final wording.  The indent is applied directly rather than through a list
# style, because the template's list styles carry their own numbering and would
# produce "1. O1  Curate ...".
# ---------------------------------------------------------------------------
log("\n--- Chapter 1: labelling the objectives and deliverables ---")
n_lab = 0
for _prefix, _label in NUMBERED_LISTS:
    _hit = None
    for _p in doc.paragraphs:
        _t = para_text(_p).strip()
        if _t.startswith(_prefix):
            _hit = _p
            break
    if _hit is None:
        raise SystemExit("Chapter 1 list item not found: " + _prefix)
    rewrite_paragraph_text(_hit, _label + " " + para_text(_hit).strip())
    _pf = _hit.paragraph_format
    _pf.left_indent = Inches(0.5)
    _pf.first_line_indent = Inches(-0.35)
    n_lab += 1
log(f"  {n_lab} item(s) labelled O1-O4 / D1-D4 and indented")


# ---------------------------------------------------------------------------
# Chapter 3 -- build the blocks and move each one before the anchor
# ---------------------------------------------------------------------------
log("\n--- Chapter 3: inserting blocks ---")
STYLE_FOR = {"h1": "1", "h2": "2", "h3": "3",
             "figurecaption": "figurecaption", "tablecaption": "tablecaption"}
counts = {}
placeholders = []


def emit(el):
    """Move a freshly appended element to just before the anchor heading."""
    ref_heading._p.addprevious(el)


def build_block(block):
    """Render one content block and return the elements it produced.

    The elements are created at the end of the document and are not yet in
    their final position; the caller decides where they go.  Chapters 3 to
    6 are appended before the references heading, while Chapter 2's
    additions are spliced in after an anchor paragraph inside prose that
    already exists -- two different destinations needing one shared
    vocabulary of blocks.
    """
    made = []
    kind = block[0]
    counts[kind] = counts.get(kind, 0) + 1

    if kind in ("h1", "h2", "h3"):
        p = doc.add_paragraph(block[1])
        set_style_id(p, STYLE_FOR[kind])
        made.append(p._p)

    elif kind == "p":
        p = doc.add_paragraph(block[1])
        set_style_id(p, "a")
        made.append(p._p)

    elif kind in ("figurecaption", "tablecaption"):
        p = doc.add_paragraph(block[1])
        set_style_id(p, STYLE_FOR[kind])
        made.append(p._p)

    elif kind == "eq":
        p = doc.add_paragraph()
        set_style_id(p, "a")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(block[1])
        r.italic = True
        made.append(p._p)

    elif kind == "image":
        p = doc.add_paragraph()
        set_style_id(p, "a")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if block[1].startswith("[["):
            # an asset only the author can supply; leave a visible marker
            # rather than a silent gap, and log it for the pending register
            r = p.add_run(block[1])
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            placeholders.append(block[1])
        else:
            path = os.path.join(FIG, block[1])
            if not os.path.exists(path):
                raise SystemExit(f"missing figure: {path}")
            p.add_run().add_picture(path, width=Inches(block[2]))
        made.append(p._p)

    elif kind == "table":
        rows = block[1]
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = doc.styles["Table Grid"]
        t.autofit = True
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                c = t.cell(i, j)
                c.text = ""
                par = c.paragraphs[0]
                set_style_id(par, "a")
                run = par.add_run(str(cell))
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
        made.append(t._tbl)

    elif kind == "algo":
        cap = doc.add_paragraph(block[1])
        set_style_id(cap, "tablecaption")
        made.append(cap._p)
        t = doc.add_table(rows=1, cols=1)
        t.style = doc.styles["Table Grid"]
        cell = t.cell(0, 0)
        cell.text = ""
        first = True
        for line in block[2]:
            par = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            set_style_id(par, "a")
            pf = par.paragraph_format
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
            run = par.add_run(line if line else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        made.append(t._tbl)

    else:
        raise SystemExit(f"unknown block kind: {kind}")

    return made


for block in (list(CH3_BLOCKS) + list(CH4_BLOCKS) + list(CH5_BLOCKS)
              + list(CH6_BLOCKS)):
    for _el in build_block(block):
        emit(_el)

log("  blocks inserted: " + ", ".join(f"{k}={v}" for k, v in sorted(counts.items())))
# ---------------------------------------------------------------------------
# Chapter 2 -- expand the literature review in place
#
# Chapter 2 lives in the source document rather than being generated, so its
# expansion is spliced into prose that already exists.  Anchors are matched
# only against paragraphs that follow the "Chapter 1" heading: the table of
# contents repeats every heading verbatim, and an unrestricted search for
# "2.4 Choice of Methods" finds the contents entry first and would insert two
# new sections into the contents page.
# ---------------------------------------------------------------------------
log("\n--- Chapter 2: expanding the literature review ---")


def body_start_index():
    """Index of the first paragraph belonging to the body, not the preamble."""
    for i, par in enumerate(doc.paragraphs):
        if style_id_of(par) == "1" and para_text(par).startswith("Chapter 1"):
            return i
    raise SystemExit("could not locate the 'Chapter 1' heading")


def style_id_of(par):
    pPr = par._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    el = pPr.find(qn("w:pStyle"))
    return None if el is None else el.get(qn("w:val"))


def find_body_paragraph(fragment):
    """The first body paragraph containing `fragment`."""
    for par in doc.paragraphs[body_start_index():]:
        if fragment in para_text(par):
            return par
    raise SystemExit("Chapter 2 anchor not found: " + fragment[:70])


n_rep2 = 0
for _old, _rep in CH2_REPLACEMENTS:
    for _par in doc.paragraphs[body_start_index():]:
        _t = para_text(_par)
        if _old not in _t:
            continue
        rewrite_paragraph_text(_par, _t.replace(_old, _rep))
        n_rep2 += 1
log(f"  {n_rep2} paragraph(s) rewritten")

n_blocks2 = 0
for _where, _anchor, _blocks in CH2_INSERTIONS:
    _target = find_body_paragraph(_anchor)
    _elements = []
    for _b in _blocks:
        _elements.extend(build_block(_b))
    if _where == "after":
        _cur = _target._p
        for _el in _elements:
            _cur.addnext(_el)
            _cur = _el
    elif _where == "before":
        for _el in _elements:
            _target._p.addprevious(_el)
    else:
        raise SystemExit("unknown insertion position: " + _where)
    n_blocks2 += len(_blocks)
    log(f"  {len(_blocks):2d} block(s) {_where:6s} '{_anchor[:46]}...'")
log(f"  {n_blocks2} block(s) added to Chapter 2")

# Signposting paragraphs the contents page already covers.  Runs after the
# insertions above, because one of them anchors on the paragraph being removed;
# the blocks it inserted stay where they were placed.
n_del2 = 0
for _frag in CH2_DELETE:
    _hits = [_p for _p in doc.paragraphs[body_start_index():]
             if _frag in para_text(_p)]
    if len(_hits) != 1:
        raise SystemExit("Chapter 2 paragraph to delete matched %d times: %s"
                         % (len(_hits), _frag[:60]))
    _hits[0]._p.getparent().remove(_hits[0]._p)
    n_del2 += 1
    log("  removed signposting paragraph: " + _frag[:52] + "...")
log("  " + str(n_del2) + " paragraph(s) removed from Chapter 2")

# Extra rows for Table 2.1, so that the newly discussed systems appear in the
# summary alongside those already listed.
_t21 = None
for _tbl in doc.tables:
    if _tbl.rows and _tbl.cell(0, 0).text.strip() == "Prior work":
        _t21 = _tbl
        break
if _t21 is None:
    raise SystemExit("Table 2.1 not found (no table headed 'Prior work')")
for _row in TABLE_2_1_ADDITIONS:
    _cells = _t21.add_row().cells
    for _j, _text in enumerate(_row):
        _cells[_j].text = ""
        _par = _cells[_j].paragraphs[0]
        set_style_id(_par, "a")
        _run = _par.add_run(_text)
        _run.font.size = Pt(9)
log(f"  {len(TABLE_2_1_ADDITIONS)} row(s) appended to Table 2.1 "
    f"(now {len(_t21.rows)} rows)")

if placeholders:
    log(f"  {len(placeholders)} placeholder(s) awaiting an author-supplied asset:")
    for ph in placeholders:
        log(f"    {ph}")


# ---------------------------------------------------------------------------
# References -- append [28]..[50] in the style of the existing 27 entries
# ---------------------------------------------------------------------------
log("\n--- References ---")
last_ref = None
seen_refs = False
for p in doc.paragraphs:
    t = para_text(p).strip()
    if t == "List of References":
        seen_refs = True
        continue
    if seen_refs:
        if t.startswith("[") and "]" in t:
            last_ref = p
        elif t.startswith("Appendix"):
            break
if last_ref is None:
    raise SystemExit("could not find the end of the reference list")
log(f"  last existing entry: {para_text(last_ref)[:44]}...")

# Chapter 2 introduces [28]-[40] and Chapter 3 [41]-[50]; the list has to be
# emitted in numeric order, which is also order of first appearance.
ALL_REFERENCES = sorted(list(CH2_REFERENCES) + list(NEW_REFERENCES),
                        key=lambda r: r[0])
_nums = [r[0] for r in ALL_REFERENCES]
if _nums != list(range(_nums[0], _nums[0] + len(_nums))):
    raise SystemExit("new reference numbers are not consecutive: %s" % _nums)

anchor = last_ref
for num, text in ALL_REFERENCES:
    p = doc.add_paragraph(f"[{num}]  {text}")
    set_style_id(p, "a")
    anchor._p.addnext(p._p)
    anchor = p
log(f"  appended [{ALL_REFERENCES[0][0]}]..[{ALL_REFERENCES[-1][0]}] "
    f"({len(ALL_REFERENCES)} entries)")


# ---------------------------------------------------------------------------
# Front matter -- replace the Leeds template's own text with the report's
#
# The template marks its gaps with "<...>" rather than our "[[...]]", and some
# of its notes run across paragraph boundaries, so neither the placeholder
# check nor a per-paragraph regular expression saw them.  The Summary was still
# a template instruction as late as 19 August 2026.
#
# Everything here is confined to paragraphs before the Chapter 1 heading, so a
# fragment that also occurs in the body cannot be caught by accident.
# ---------------------------------------------------------------------------
log("\n--- Front matter: title page, Summary, Acknowledgements ---")


def front_matter_paragraphs():
    out = []
    for par in doc.paragraphs:
        if para_text(par).strip().startswith("Chapter 1"):
            break
        out.append(par)
    return out


n_fm = 0
for _old, _new in FM_REPLACEMENTS:
    for _p in front_matter_paragraphs():
        _t = para_text(_p)
        if _old in _t:
            rewrite_paragraph_text(_p, _t.replace(_old, _new))
            n_fm += 1
            log("  filled in: " + _new[:56])
log("  " + str(n_fm) + " placeholder(s) filled")

n_del = 0
for _frag in FM_DELETE:
    for _p in front_matter_paragraphs():
        if _frag in para_text(_p):
            _p._p.getparent().remove(_p._p)
            n_del += 1
            log("  removed template note: " + _frag[:52] + "...")
            break
    else:
        raise SystemExit("front-matter paragraph to delete not found: " + _frag)
log("  " + str(n_del) + " template instruction paragraph(s) removed")

for _frag, _blocks in FM_BLOCKS:
    _target = None
    for _p in front_matter_paragraphs():
        if _frag in para_text(_p):
            _target = _p
            break
    if _target is None:
        raise SystemExit("front-matter placeholder not found: " + _frag)
    _cur = _target._p
    for _b in _blocks:
        for _el in build_block(_b):
            _cur.addnext(_el)
            _cur = _el
    _target._p.getparent().remove(_target._p)
    log("  wrote " + str(len(_blocks)) + " paragraph(s) over: " + _frag[:44] + "...")

# The deliverables table still promised an envelope of signed consent forms.
_dtable = None
for _tbl in doc.tables:
    if _tbl.rows and _tbl.cell(0, 0).text.strip() == DELIVERABLES_HEADER:
        _dtable = _tbl
        break
if _dtable is None:
    raise SystemExit("deliverables table not found on the title page")
_was = len(_dtable.rows)
while len(_dtable.rows) > len(DELIVERABLES):
    _dtable._tbl.remove(_dtable.rows[-1]._tr)
while len(_dtable.rows) < len(DELIVERABLES):
    _dtable.add_row()
for _i, _row in enumerate(DELIVERABLES):
    _cells = _dtable.rows[_i].cells
    for _j, _text in enumerate(_row):
        _cells[_j].text = ""
        _par = _cells[_j].paragraphs[0]
        _run = _par.add_run(_text)
        _run.font.size = Pt(10)
        if _i == 0:
            _run.bold = True
log("  deliverables table rebuilt: " + str(_was) + " rows -> "
    + str(len(_dtable.rows)) + " (consent-forms row removed)")

for _p in front_matter_paragraphs():
    _t = para_text(_p)
    if "<" in _t and ">" in _t:
        raise SystemExit("a template placeholder survived: " + _t[:80])


# ---------------------------------------------------------------------------
# Appendices -- replace the template placeholder text
#
# Both appendices still carried the Leeds template's filler.  That is not
# merely untidy: Section 1.4 already tells the reader that Appendix A records
# how the corpus may be obtained, so the document was making a forward
# reference to a page that said nothing.  verify_thesis.py could not catch it,
# because its placeholder check looks for "[[" and template filler has none.
# ---------------------------------------------------------------------------
log("\n--- Appendices: replacing the template placeholder text ---")

_body = doc.paragraphs[body_start_index():]
_heading_at = {}
for _i, _par in enumerate(_body):
    _t = para_text(_par).strip()
    for _h in REPLACE_FROM_HEADING:
        if style_id_of(_par) == "1" and _t.startswith(_h):
            _heading_at.setdefault(_h, _i)
for _h in REPLACE_FROM_HEADING:
    if _h not in _heading_at:
        raise SystemExit("appendix heading not found: " + _h)

for _h, _blocks in (("Appendix A", APPENDIX_A), ("Appendix B", APPENDIX_B)):
    _start = _heading_at[_h]
    # Everything from just after this heading to the next appendix heading (or
    # the end of the body) is placeholder text and goes.
    _stop = len(_body)
    for _other, _idx in _heading_at.items():
        if _idx > _start:
            _stop = min(_stop, _idx)
    _removed = 0
    for _par in _body[_start + 1:_stop]:
        _par._p.getparent().remove(_par._p)
        _removed += 1

    _anchor_par = _body[_start]
    _cur = _anchor_par._p
    for _b in _blocks:
        for _el in build_block(_b):
            _cur.addnext(_el)
            _cur = _el
    log(f"  {_h}: removed {_removed} placeholder paragraph(s), "
        f"inserted {len(_blocks)} block(s)")

for _par in doc.paragraphs:
    if "Text under appendix heading" in para_text(_par) \
            or "Text under level" in para_text(_par):
        raise SystemExit("template placeholder text survived in the appendices")


# ---------------------------------------------------------------------------
# Captions -- set every one to body size
#
# The supervisor asked on 20 August 2026 that "the font size for caption for
# figures and tables should be the same size as text".  Most captions already
# inherited it: only the four that came with Chapter 2 in the source document
# carried a hard 9pt run, which is why the two he happened to mark were both
# in that chapter.  Setting the size explicitly on every caption makes the
# document uniform whatever a future caption inherits.
# ---------------------------------------------------------------------------
log("\n--- Captions: set to body size ---")
_body_pt = doc.styles["Normal"].font.size
n_cap = 0
for _p in doc.paragraphs:
    if _p.style.name not in ("figure caption", "table caption"):
        continue
    for _r in _p.runs:
        if _r.font.size != _body_pt:
            _r.font.size = _body_pt
            n_cap += 1
log(f"  {n_cap} caption run(s) set to {_body_pt.pt:g}pt")


doc.save(DST)
shutil.copyfile(DST, LATEST)
log("")
log(f"saved: {os.path.relpath(DST, ROOT)}   (versioned, for rollback)")
log(f"       {os.path.relpath(LATEST, ROOT)}   <-- OPEN THIS ONE")
log("NOTE: open in Word and press Ctrl+A then F9 to rebuild the "
    "table of contents.")

with open(os.path.join(HERE, "outputs", "build_report.txt"), "w",
          encoding="utf-8") as f:
    f.write("\n".join(report))
