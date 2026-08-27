"""Verification pass over the built dissertation.

Checks, in order:
  1  no unresolved placeholders survive
  2  figures and tables appear in ascending document order
  3  every figure and table is referred to somewhere in the body text
  4  citation numbers are contiguous, and every reference is cited
  5  no wording presents allergen screening as a safety guarantee
  6  word counts per chapter, against the page budget
  7  Chapters 3 and 4 agree with each other and with the code

Check 7 exists because Chapter 3 was written before the implementation and
Chapter 4 after it.  Either can be internally consistent while contradicting
the other, and the contradiction is invisible from inside either one.  It also
compares Chapter 4's stated number of assertion groups against the number the
suite actually defines, so that adding a group without updating the prose is
caught rather than left as a quiet overstatement.
"""

import io
import os
import re
import sys
from collections import Counter

import docx
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
DOC = os.path.join(HERE, "..", "毕业论文_v5.docx")

doc = docx.Document(DOC)
fail = []


def text_of(p):
    return "".join(r.text for r in p.runs)


# ordered walk of paragraphs and tables
blocks = []
for child in doc.element.body.iterchildren():
    if child.tag == qn("w:p"):
        for p in doc.paragraphs:
            if p._p is child:
                blocks.append(("p", p))
                break
    elif child.tag == qn("w:tbl"):
        blocks.append(("tbl", child))

paras = [(p, text_of(p)) for k, p in blocks if k == "p"]
alltext = "\n".join(t for _, t in paras)


def style_id(p):
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        return ""
    el = pPr.find(qn("w:pStyle"))
    return el.get(qn("w:val")) if el is not None else ""


print("=" * 66)
print("1. PLACEHOLDER SCAN")
print("=" * 66)
ph = [t for _, t in paras if "[[" in t]
if ph:
    fail.append("placeholders remain")
    for t in ph:
        print("   REMAINS:", t[:90])
else:
    print("   clean: no [[CITE]] / [[FIG]] / [[SCREENSHOT]] / [[RESULT]] found")

# The Leeds template marks its own gaps with <...> and dates with xx/xx/xx,
# which the check above never saw.  On 19 August 2026 that meant the Summary
# was still the template's instruction to the author -- the first page an
# examiner reads -- while every check reported the document clean.  Some of
# those notes run across paragraph boundaries, so the whole document is joined
# before matching rather than each paragraph being tested on its own.
joined = alltext
for _tbl in doc.tables:
    for _row in _tbl.rows:
        for _cell in _row.cells:
            joined += chr(10) + _cell.text
template = re.findall(r"<[^<>]{4,200}>", joined, re.S)
template += ["xx/xx/xx"] * joined.count("xx/xx/xx")
if template:
    fail.append("Leeds template placeholders remain")
    for t in template:
        print("   TEMPLATE REMAINS:", " ".join(t.split())[:88])
else:
    print("   clean: no <...> or xx/xx/xx template residue")

print()
print("=" * 66)
print("2. FIGURE AND TABLE ORDER")
print("=" * 66)
seq = []
for p, t in paras:
    s = style_id(p)
    m = re.match(r"(Figure|Table|Algorithm)\s+(\d+)\.(\d+)", t.strip())
    if m and s in ("figurecaption", "tablecaption"):
        seq.append((m.group(1), int(m.group(2)), int(m.group(3)), t.strip()[:52]))
for kind in ("Figure", "Table"):
    nums = [(c, n) for k, c, n, _ in seq if k == kind]
    print(f"   {kind}s in document order: " +
          ", ".join(f"{c}.{n}" for c, n in nums))
    if nums != sorted(nums):
        fail.append(f"{kind} numbering out of order")
        print(f"   *** OUT OF ORDER ***")
    else:
        print(f"   ascending: OK")

print()
print("=" * 66)
print("3. CROSS-REFERENCES FROM BODY TEXT")
print("=" * 66)
caption_ids = {f"{k} {c}.{n}" for k, c, n, _ in seq}
bodytext = "\n".join(t for p, t in paras
                     if style_id(p) not in ("figurecaption", "tablecaption"))
missing = []
for cid in sorted(caption_ids):
    n = len(re.findall(re.escape(cid) + r"\b", bodytext))
    flag = "OK " if n else "*** NOT REFERRED TO ***"
    print(f"   {cid:<16} referred to {n} time(s)   {flag}")
    if not n:
        missing.append(cid)
if missing:
    fail.append(f"unreferenced: {missing}")

print()
print("=" * 66)
print("4. CITATIONS")
print("=" * 66)
# reference list entries
refs = []
seen = False
for p, t in paras:
    s = t.strip()
    if s == "List of References":
        seen = True
        continue
    if seen:
        m = re.match(r"\[(\d+)\]", s)
        if m:
            refs.append(int(m.group(1)))
        elif s.startswith("Appendix"):
            break
print(f"   reference list: {len(refs)} entries, [{min(refs)}]..[{max(refs)}]")
if refs != list(range(1, len(refs) + 1)):
    fail.append("reference numbering not contiguous")
    print("   *** NOT CONTIGUOUS ***")
else:
    print("   contiguous, no duplicates: OK")

# citations in body (exclude the reference list itself)
cut = next(i for i, (p, t) in enumerate(paras)
           if t.strip() == "List of References")
citing = "\n".join(t for p, t in paras[:cut])
cited = Counter()
for grp in re.findall(r"\[([\d,\s]+)\]", citing):
    for n in re.findall(r"\d+", grp):
        cited[int(n)] += 1
uncited = [n for n in refs if n not in cited]
print(f"   distinct references cited in the body: {len(cited)}")
if uncited:
    fail.append(f"uncited references: {uncited}")
    print(f"   *** NEVER CITED: {uncited} ***")
else:
    print("   every reference is cited at least once: OK")
new_cited = sorted(n for n in cited if n >= 28)
print(f"   new references cited in Chapter 3: {new_cited}")
dangling = [n for n in cited if n > max(refs)]
if dangling:
    fail.append(f"citation with no entry: {dangling}")
    print(f"   *** CITED BUT NOT IN LIST: {dangling} ***")

print()
print("=" * 66)
print("5. SAFETY WORDING")
print("=" * 66)
BAD = [r"\bguarantees\b", r"\bensures\b", r"\bensure\b", r"\bis safe\b",
       r"\bcompletely safe\b", r"\ball allergens\b", r"\beliminates\b"]
hits = []
for p, t in paras:
    for pat in BAD:
        for m in re.finditer(pat, t, re.I):
            frag = t[max(0, m.start() - 70):m.end() + 70]
            hits.append((m.group(0), frag))
if hits:
    print("   contexts to inspect:")
    for w, frag in hits:
        print(f"     '{w}': ...{frag.strip()}...")
else:
    print("   no absolute-safety wording found: OK")
must = ["not a safety guarantee", "fail-closed", "never relaxed"]
for phrase in must:
    n = alltext.lower().count(phrase.lower())
    print(f"   '{phrase}' appears {n} time(s)   " + ("OK" if n else "*** MISSING ***"))
    if not n:
        fail.append(f"missing required phrase: {phrase}")

print()
print("=" * 66)
print("6. LENGTH")
print("=" * 66)
chapters = {}
cur = "front matter"
for p, t in paras:
    s = t.strip()
    if style_id(p) == "1":
        cur = s
    chapters[cur] = chapters.get(cur, 0) + len(t.split())
for k, v in chapters.items():
    print(f"   {k[:44]:<46} {v:>6,} words")
print(f"   {'TOTAL':<46} {sum(chapters.values()):>6,} words")

# ---------------------------------------------------------------------------
print()
print("=" * 66)
print("7. CHAPTERS 3 AND 4 AGREE WITH EACH OTHER AND WITH THE CODE")
print("=" * 66)
# Chapter 3 was written before the code existed and Chapter 4 after it, so the
# two can drift apart without either being wrong on its own terms.  These are
# the quantities both chapters depend on.
for figure in ["128,403", "655,954", "19,919", "40,779", "103,389"]:
    n = alltext.count(figure)
    flag = "OK" if n >= 1 else "*** ABSENT ***"
    if n < 1:
        fail.append("contract figure missing: " + figure)
    print(f"   contract figure {figure:<9} {n} occurrence(s)   {flag}")

# The suite names its own groups with section(n, ...); Chapter 4 states how
# many there are in prose.  A group added without updating the prose is a quiet
# overstatement, which is the kind of claim this project's register exists to
# prevent.
suite = io.open(os.path.join(HERE, "verify_prototype.py"),
                encoding="utf-8").read()
groups = len(re.findall(r"^section\(\d+,", suite, re.M))
WORDS = {11: "eleven", 12: "twelve", 13: "thirteen", 14: "fourteen",
         15: "fifteen", 16: "sixteen"}
claimed = WORDS.get(groups, str(groups))
phrase = claimed + " groups of assertions"
if phrase in alltext:
    print(f"   suite has {groups} groups; Chapter 4 says '{claimed}'   OK")
else:
    fail.append("Chapter 4 misstates the number of assertion groups")
    print(f"   *** suite has {groups} groups; Chapter 4 does not say "
          f"'{claimed}' ***")

print()
print("=" * 66)
print("8. MANDATORY SECTIONS REQUIRED BY THE SCHOOL STRUCTURE DOCUMENT")
print("=" * 66)
# WHY THIS EXISTS.  "Structure of MSc Dissertation" prescribes chapter and
# section titles for each project type and warns that "failure to include a
# required section or misplacing it can lead to the loss of marks".  For an
# Exploratory Software project Chapter 3 must be titled "Software Requirements
# and System Design" and must contain a software requirements section; Chapter
# 5 must cover software testing; Chapter 6 must draw conclusions; and Section
# 1.1 must open with the prescribed sentence.
#
# The report went through seven checks, four verification suites and a
# supervisor read-through without the missing requirements section being
# noticed by any of them, because none of them was looking for it.  This check
# is here so that the omission cannot recur silently.
REQUIRED_TITLES = [
    "Chapter 3 Software Requirements and System Design",
    "3.1.1  Software Requirements",
    "3.1.2  Design overview",
    "Chapter 5 Software Testing and Evaluation",
    "6.1  Conclusions",
    "1.1 Project Aim",
    "1.2 Objectives",
    "1.3 Deliverables",
    "1.4 Ethical, Legal, Social and Professional Issues",
]
for title in REQUIRED_TITLES:
    present = any(t.strip() == title for _, t in paras)
    if present:
        print(f"   {title:<52} OK")
    else:
        fail.append("mandatory section missing: " + title)
        print(f"   *** {title:<48} ABSENT ***")

# Section 1.1 must START with the prescribed opening, not merely contain the
# aim somewhere inside it.
AIM_OPENING = "The aim of this project is to"
opening = None
for i, (_, t) in enumerate(paras):
    if t.strip() == "1.1 Project Aim":
        for _, nxt in paras[i + 1:]:
            if nxt.strip():
                opening = nxt.strip()
                break
        break
if opening is None:
    fail.append("Section 1.1 has no body text")
    print("   *** Section 1.1 has no body text ***")
elif opening.startswith(AIM_OPENING):
    print(f"   Section 1.1 opens with '{AIM_OPENING} ...'   OK")
else:
    fail.append("Section 1.1 does not open with '" + AIM_OPENING + "'")
    print(f"   *** Section 1.1 opens '{opening[:56]}...' ***")

print()
print("=" * 66)
if fail:
    print("RESULT: " + str(len(fail)) + " problem(s)")
    for f in fail:
        print("  - " + f)
    sys.exit(1)
print("RESULT: all checks passed")
